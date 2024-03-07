import sys, os
import re 
from docx import Document
import csv
import openpyxl
from pathlib import Path

report = Document()

files_patterns = [r'^Сведения о всех типах закупок [0-9]{2}\.[0-9]{2}\.[0-9]{4}', r'^ContractSearch\(.*\)_[0-9]{2}\.[0-9]{2}\.[0-9]{4}']
eb_pattern = [r'^PrintScroller_[0-9]{2}\-[0-9]{2}\-[0-9]{4} [0-9]{2}\-[0-9]{2}']
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

class BerezkaContract():
    def __init__(self, contract_number, contract_date, contract_price, contract_customer):
        self.contract_number = contract_number
        self.contract_date = contract_date
        self.contract_price = contract_price
        self.contract_customer = contract_customer
    def __hash__(self):
        return hash(('num', self.contract_number))
    def __eq__(self, other):
        return self.contract_number==other.contract_number

class EisContract():
    def __init__(self, contract_number, contract_date, contract_reestr_number, contract_customer):
        self.contract_number = contract_number
        self.contract_date = contract_date
        self.contract_reestr_number = contract_reestr_number
        self.contract_customer = contract_customer
    def __hash__(self):
        return hash(('num', self.contract_number))
    def __eq__(self, other):
        return self.contract_number==other.contract_number

class EbContract():
    def __init__(self, contract_number, contract_date, contract_price, contract_customer):
        self.contract_number = contract_number
        self.contract_date = contract_date
        #self.contract_reestr_number = contract_reestr_number
        self.contract_price = contract_price
        self.contract_customer = contract_customer
    def __hash__(self):
        return hash(('num', self.contract_number))
    def __eq__(self, other):
        return self.contract_number==other.contract_number

def run(inn):
    temp_folder = resource_path(str(Path.home() / "Downloads"))

    berezka_contracts = []
    eis_contracts = []
    eb_contracts = []
    paths = []

    for file_name in os.listdir(path=temp_folder):
        if os.path.isfile(temp_folder + "\\" + file_name):
            if  re.compile(files_patterns[0]).search(file_name) is not None:
                paths.append(temp_folder + "\\" + file_name)
                wb = openpyxl.load_workbook(filename = f'{temp_folder}\\{file_name}')
                sheet = wb.worksheets[0]
                for index in range(8, len(sheet['A'])):
                    berezka_contracts.append(BerezkaContract(sheet['A'][index].value, sheet['E'][index].value[:-6], sheet['L'][index].value, sheet['M'][index].value))
                    #print((sheet['A'][index].value, sheet['E'][index].value, sheet['L'][index].value, sheet['M'][index].value))
                os.remove(f'{temp_folder}\\{file_name}')
            elif re.compile(files_patterns[1]).search(file_name) is not None:
                #print('2')
                with open(f'{temp_folder}\\{file_name}') as csvfile:
                    paths.append(temp_folder + "\\" + file_name)
                    cvsreader = csv.DictReader(csvfile, delimiter=';')
                    for row in cvsreader:
                        eis_contracts.append(EisContract(row['Контракт: номер'], row['Контракт: дата'], 
                        row['Цена контракта'], row['Заказчик: наименование']))
                        #row['Номер реестровой записи контракта'], row['Заказчик: наименование']))
                        #print(row['Контракт: номер'], row['Контракт: дата'], 
                        #row['Номер реестровой записи контракта'], row['Заказчик: наименование'])
                os.remove(f'{temp_folder}\\{file_name}')   
            elif re.compile(eb_pattern[0]).search(file_name) is not None:
                #print('3')
                paths.append(temp_folder + "\\" + file_name)
                work_columns = []
                wb = openpyxl.load_workbook(filename = f'{temp_folder}\\{file_name}')
                sheet = wb.worksheets[0]
                # ws = wb.worksheets[0]
                bo_do = 'none'
                bo_cols = ["Номер документа основания", "Дата документа-основания", "Сумма в валюте РФ", "Наименование получателя бюджетных средств"]
                do_cols = ["Номер подтверждающего документа", "Дата подтверждающего документа", "Сумма в валюте обязательства", "Наименование получателя БС"]
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value == "Сведения о бюджетном обязательстве/об изменении бюджетного обязательства":
                            bo_do = 'bo'
                        elif cell.value == "Сведения о денежных обязательствах":
                            bo_do = 'do'
                        if bo_do == 'bo':
                            #if cell.value == "Номер документа основания" or cell.value == "Дата документа-основания" or cell.value == "УНРЗ в реестре контрактов/реестре соглашений" or cell.value == "Сумма в валюте РФ" or cell.value == "Наименование получателя бюджетных средств": 
                            if cell.value in bo_cols:    
                                work_columns.insert(bo_cols.index(cell.value), cell.column_letter)
                        if bo_do == 'do':
                            #if cell.value == "Номер подтверждающего документа" or cell.value == "Дата подтверждающего документа" or cell.value == "Учетный номер ДО" or cell.value == "Сумма в валюте обязательства" or cell.value == "Наименование получателя БС": 
                            if cell.value in do_cols:    
                                work_columns.insert(do_cols.index(cell.value), cell.column_letter)
                print(work_columns)
                for index in range (5, len(sheet['A'])):
                    eb_contracts.append(EbContract(sheet[work_columns[0]][index].value, sheet[str(work_columns[1])][index].value, sheet[str(work_columns[2])][index].value, sheet[str(work_columns[3])][index].value))
                    #eb_contracts.append(EbContract(sheet[index][str(work_columns[0])].value, sheet[index][str(work_columns[1])].value, sheet[index][str(work_columns[2])].value, sheet[index][str(work_columns[3])].value, sheet[index][str(work_columns[4])].value))
                    #except: pass
                os.remove(f'{temp_folder}\\{file_name}')
                # sheet = wb.worksheets[0]
                # for index in range(5, len(sheet['A'])):
                #     eb_contracts.append(EbContract(sheet['T'][index].value, sheet['AP'][index].value, sheet['AT'][index].value, sheet['V'][index].value, sheet['O'][index].value))
                    #rint(sheet['T'][index].value, sheet['AP'][index].value, sheet['AT'][index].value, sheet['V'][index].value, sheet['O'][index].value)

    # top = Element('Contracts')
    # berezka = SubElement(top,'Berezka')
    # eis = SubElement(top,'EIS')

    # if len(berezka_contracts)!=0:
    #     for item in berezka_contracts:
    #         berezka_elem = SubElement(berezka, f'Contract')
    #         SubElement(berezka_elem, 'contract_number').text = item.contract_number
    #         SubElement(berezka_elem, 'contract_date').text = item.contract_date
    #         SubElement(berezka_elem, 'contract_price').text = item.contract_price
    #         SubElement(berezka_elem, 'contract_customer').text = item.contract_customer

    # if len(eis_contracts)!=0:
    #     prev_contract = ''
    #     for item in eis_contracts:
    #         if prev_contract == item.contract_reestr_number:
    #             pass
    #         else:
    #             prev_contract = item.contract_reestr_number
    #             eis_elem = SubElement(eis, f'Contract')
    #             SubElement(eis_elem, 'contract_number').text = item.contract_number
    #             SubElement(eis_elem, 'contract_date').text = item.contract_date
    #             SubElement(eis_elem, 'contract_reestr_number').text = item.contract_reestr_number
    #             SubElement(eis_elem, 'contract_customer').text = item.contract_customer
    eat_ = 0
    eis_ = 0
    missing_contrs = []
    found_contrs = []
    double_check = []
    # for i in eb_contracts:
    #     for j in eis_contracts:
    #         if i.contract_number == j.contract_number:
    #             eis += 1
    #         else: pass
    #     for j in berezka_contracts:
    #         if i.contract_number == j.contract_number:
    #             eat += 1
    #         else: pass
    for eis in eis_contracts:
        for eb in eb_contracts:
            if eis.contract_number == eb.contract_number:
                eis_ += 1
                found_contrs.append(eb)
            else: pass
        if eis not in found_contrs:
            missing_contrs.append(eis)
    for berezka in berezka_contracts:
        for eb in eb_contracts:
            if berezka.contract_number == eb.contract_number:
                eat_ += 1
                found_contrs.append(eb)
            else: pass
        if berezka not in found_contrs:
            missing_contrs.append(berezka)
    print(eat_, eis_)
    for i in berezka_contracts:
        print('B - ',i.contract_number, i.contract_date, i.contract_price, i.contract_customer)
    for i in eis_contracts:
        print('E - ',i.contract_number, i.contract_date, i.contract_reestr_number, i.contract_customer)
    for i in eb_contracts:
        print('EB - ', i.contract_number, i.contract_date, i.contract_price, i.contract_customer)
    print('missing -',list(set(missing_contrs)))
    print('found -',list(set(found_contrs)))
    table = report.add_table(1, 4)
    table.style = 'Table Grid'  
    head_cells = table.rows[0].cells
    for i, item in enumerate(['№ Документа-основания', 'Дата документа-основания', 'Сумма', 'Заказчик']):
        p = head_cells[i].paragraphs[0]
        # название колонки
        p.add_run(item).bold = True
    # добавляем данные к существующей таблице
    for i in list(set(missing_contrs)):
        # добавляем строку с ячейками к объекту таблицы
        cells = table.add_row().cells
        try:
            for n, item in enumerate([i.contract_number, i.contract_date, i.contract_price, i.contract_customer]):
                # вставляем данные в ячейки
                cells[n].text = str(item)
        except:
            for n, item in enumerate([i.contract_number, i.contract_date, i.contract_reestr_number, i.contract_customer]):
                # вставляем данные в ячейки
                cells[n].text = str(item)
    report.save(f'Отчет по {inn}.docx')
    # for i in eis_contracts:
    #     if i in eb_contracts:
    #         print('yippy')
    #     else:
    #         print('nyippy')

    # for i in berezka_contracts:
    #     if i in eb_contracts:
    #         print('yay')
    #     else:
    #         print('nan')

    # if len(eis_contracts)==0 and len(berezka_contracts)==0:
    #     print('000000sf')
    # else:
    #     xml_filename = f'{datetime.now().strftime("%m%d%Y%H%M%S")}.xml'
    #     ElementTree(top).write(f'{temp_folder}\\{xml_filename}')

        #save_file.emit(xml_filename)

    # for i in berezka_contracts:
    #     os.remove(i)
    # for i in eis_contracts:
    #     os.remove(i)

#run('7709895509')
